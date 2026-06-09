from docx import Document
from pathlib import Path


def parse_docx(file_path: str) -> str:
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        text = f"Word解析错误: {str(e)}"
    return text
